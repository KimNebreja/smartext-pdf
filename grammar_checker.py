import re
import os
import time
import json
import asyncio
from typing import List
from dotenv import load_dotenv

import win32com.client
from docx import Document
from docx2pdf import convert
import requests

from openai import AsyncOpenAI
from fastapi import FastAPI, HTTPException, Query
from fastapi.responses import JSONResponse
import uvicorn

# Load environment variables from .env file for API keys, etc.
load_dotenv()
api_key = os.getenv("API_KEY")

# Initialize OpenAI client for GPT-based grammar correction
client = AsyncOpenAI(api_key=api_key)

#Download original file from Hostinger


def download_from_hostinger(file_code, save_path):
    url = f"https://smartext-pdf.com/original_pdfs/{file_code}.pdf"
    response = requests.get(url)
    if response.status_code == 200:
        with open(save_path, "wb") as f:
            f.write(response.content)
        print(f"Downloaded {file_code}.pdf from Hostinger")
    else:
        raise Exception(f"Failed to download {file_code}.pdf from Hostinger")
    
def download_processed_from_hostinger(file_code, dst_pdf):
    url = f"https://smartext-pdf.com/processed_pdfs/{file_code}.pdf"
    r = requests.get(url)
    if r.status_code == 200:
        with open(dst_pdf, "wb") as f:
            f.write(r.content)
    else:
        raise Exception(f"Failed to fetch processed PDF {file_code} from Hostinger")

def download_json_from_hostinger(file_code, dst_json):
    url = f"https://smartext-pdf.com/jsons/{file_code}.json"
    r = requests.get(url)
    if r.status_code == 200:
        with open(dst_json, "wb") as f:
            f.write(r.content)
    else:
        raise Exception(f"Failed to fetch JSON {file_code} from Hostinger")



#Upload processed file to Hostinger
   
def upload_processed_file(final_pdf_path, file_code):
    # Full URL to your PHP file on Hostinger
    upload_url = "https://smartext-pdf.com/upload_processed.php"


    files = {
        "file": (os.path.basename(final_pdf_path), open(final_pdf_path, "rb"), "application/pdf"),
        "file_code": (None, file_code)
    }


    response = requests.post(upload_url, files=files)


    if response.status_code == 200:
        print("✅ Uploaded processed file:", response.text)
    else:
        print("❌ Upload failed:", response.status_code, response.text)




def upload_parsing_words(updated_docx_path, file_code):
    # Full URL to your PHP file on Hostinger
    upload_url = "https://smartext-pdf.com/upload_parse.php"


    files = {
        "file": (os.path.basename(updated_docx_path), open(updated_docx_path, "rb"), "application/pdf"),
        "file_code": (None, file_code)
    }


    response = requests.post(upload_url, files=files)


    if response.status_code == 200:
        print("✅ Uploaded processed file:", response.text)
    else:
        print("❌ Upload failed:", response.status_code, response.text)

def upload_jsons(json_output_path, file_code):
    # Full URL to your PHP file on Hostinger
    upload_url = "https://smartext-pdf.com/upload_jsons.php"


    files = {
        "file": (os.path.basename(json_output_path), open(json_output_path, "rb"), "application/pdf"),
        "file_code": (None, file_code)
    }


    response = requests.post(upload_url, files=files)


    if response.status_code == 200:
        print("✅ Uploaded processed file:", response.text)
    else:
        print("❌ Upload failed:", response.status_code, response.text)

# ---------------------------
# PDF to Word Conversion
# ---------------------------
def convert_pdf_to_word(pdf_path, docx_path):
    """
    Convert a PDF file to a Word (.docx) file using Microsoft Word automation.
    """
    print("Converting PDF to Word using Microsoft Word...")
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(pdf_path)
    doc.SaveAs(docx_path, FileFormat=16)
    doc.Close()
    word.Quit()
    print("Conversion done.")

# ---------------------------
# Tokenizer with punctuation
# ---------------------------
def tokenize_with_punctuation(text):
    """
    Tokenizer that splits punctuation as separate tokens.
    """
    return re.findall(r"\w+|[^\w\s]", text, re.UNICODE)

# ---------------------------
# Word to PDF Conversion
# ---------------------------
def convert_word_to_pdf(updated_docx_path, final_pdf_path):
    """
    Convert a Word (.docx) file back to PDF.
    """
    print("Converting back to PDF...")
    convert(updated_docx_path, final_pdf_path)
    print(f"Final PDF saved as: {final_pdf_path}")


# ---------------------------
# GPT-based grammar correction
# ---------------------------
async def gpt_proofread(text):
    """
    Calls OpenAI GPT using a function/tool to proofread and classify changes.
    Ensures the structured JSON format and retries up to 3 times if needed.
    """

    # Normalize excessive spaces
    text = re.sub(r"[^\S\r\n]{2,}", " ", text)

    tool_schema = {
        "name": "proofread_output",
        "description": "Returns corrected content with detailed revisions for grammar, formality, and punctuation.",
        "parameters": {
            "type": "object",
            "properties": {
                "original": {"type": "string"},
                "corrected": {"type": "string"},
                "original_token": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "idx": {"type": "integer"},
                            "word": {"type": "string"}
                        },
                        "required": ["idx", "word"]
                    }
                },
                "proofread_token": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "idx": {"type": "integer"},
                            "word": {"type": "string"}
                        },
                        "required": ["idx", "word"]
                    }
                },
                "changes": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "type": {
                                "type": "string",
                                "enum": ["replaced", "corrected", "inserted", "removed"]
                            },
                            "original_idx": {"type": ["integer", "null"]},
                            "proofread_idx": {"type": ["integer", "null"]},
                            "original_word": {"type": "string"},
                            "proofread_word": {"type": "string"},
                            "suggestion": {
                            "type": "array",
                            "items": {
                                "type": "string",
                                "pattern": "^[^\\s]+$"
                            },
                            "maxItems": 3
                            }
                        },
                        "required": [
                            "type", "original_idx", "proofread_idx",
                            "original_word", "proofread_word", "suggestion"
                        ]
                    }
                }
            },
            "required": ["original", "corrected", "original_token", "proofread_token", "changes"]
        }
    }


    system_msg = (
        "You are a professional AI assistant integrated into a proofreading system called SmarText PDF. \n"
        "You will receive raw text extracted from PDFs and must correct grammar, punctuation, and spelling errors. \n"
        "Improve the tone to be more formal and ensure the text is clear and coherent, while keeping the original meaning intact. "
        "You may receive content in various languages. Proofread in the language the text is written in — be bilingual and flexible. \n"
        "Do not translate unless explicitly asked. Your goal is to refine the original language, not convert it. \n\n"

        "Do not modify names of people, places, companies, or titles — even if they appear to be misspelled, uncapitalized, or stylized. "
        "Preserve them exactly as they appear in the original text, unless the correction is absolutely unambiguous (e.g., fixing 'jonh' to 'John' when the context clearly indicates a typo). "
        "Avoid guessing or rewriting names, especially rare, creative, or user-generated ones like handles or nicknames. \n\n"

        "Examples:\n"
        "- Leave names like 'john smith', 'Gooogle', 'mcdonalds', or 'Elonn' unchanged.\n"
        "- Do not correct 'Jhon' to 'John' unless it's clearly a typo.\n"
        "- Do not change 'may' to 'might' if 'May' is a name or month.\n"
        "- Do not alter stylized names like 'iPhone', 'eBay', or usernames like '@daniel42'.\n\n"

        "Your output must strictly use the function proofread_output with all required fields: "
        "original, corrected, original_token, proofread_token, and changes.\n\n"

        "Each token must be processed as a word or punctuation mark — no grouping or skipping. "
        "If punctuation changes (e.g., '.', ',', '?', '!') or is added/removed, you must include it in changes. "
        "If a token is part of a numbered heading or subheading (e.g., '1 ', '1.', '1.1.', 'I.', 'A)'), retain the numbering or Roman numeral as-is. Do not modify or remove it. "
        "Maintain accurate idx for both original_token and proofread_token. \n\n"

        "For each change:\n"
        "- Use corrected for fixing grammar or punctuation with minor word adjustments.\n"
        "- Use replaced when words or phrases are changed into different terms (and optionally include up to 3 suggestions).\n"
        "- Use inserted if new words were added.\n"
        "- Use removed if unnecessary words or punctuation were deleted.\n\n"

        "If multiple rewrites are possible, always provide up to 3 meaningful alternatives in the suggestions array, not just the accepted version. "
        "For example, if correcting 'fastly' to 'quickly', suggestions could be ['quickly', 'rapidly', 'swiftly']. "
        "If only one correction exists, include it alone. "
        "Every change must include original_word, proofread_word, and suggestion even if suggestions only include the accepted version. "
        "Ensure the suggestion list is meaningful — if only one correction exists, include it alone. If multiple rewrites are possible, list alternatives.\n\n"

        "Do not skip punctuation or formatting changes. The goal is accurate proofreading that can be traced token-by-token and visually rendered with detailed changes."
    )

    


    def ensure_structure(result, original_text):
        return {
            "original": result.get("original", original_text),
            "corrected": result.get("corrected", original_text),
            "original_token": result.get("original_token", []),
            "proofread_token": result.get("proofread_token", []),
            "changes": result.get("changes", [])
        }

    for attempt in range(1, 4):
        try:
            response = await client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": f"Proofread this sentence:\n{text}"}
                ],
                tools=[{"type": "function", "function": tool_schema}],
                tool_choice="auto",
                temperature=0.3,
            )

            function_call = response.choices[0].message.tool_calls[0].function
            result = json.loads(function_call.arguments)
            result = ensure_structure(result, text)

            # Check completeness
            if result["original_token"] and result["proofread_token"]:
                return result
            else:
                print(f"[WARNING] Attempt {attempt}: Missing token data.")
        except Exception as e:
            print(f"[ERROR] Attempt {attempt}: Tool call failed.\nInput: {repr(text[:200])}")
            print("Error:", e)

    # Fallback return
    print("[ERROR] All 3 attempts failed. Returning fallback structure.")
    return {
        "original": text,
        "corrected": text,
        "original_token": [],
        "proofread_token": [],
        "changes": []
    }
# ---------------------------
# Async wrapper for GPT proofreading
# ---------------------------
async def async_gpt_proofread(paragraph_id, text):
    """
    Async wrapper to call GPT proofreading for a paragraph.
    Returns paragraph_id, original text, and GPT response.
    """
    try:
        gpt_response = await gpt_proofread(text)
        return paragraph_id, text, gpt_response
    except Exception as e:
        print(f"Error processing paragraph {paragraph_id}: {e}")
        return paragraph_id, text, {"corrected": text}
    
# ---------------------------
# Main async grammar correction for all paragraphs
# ---------------------------
# ---------------------------
# Main async grammar correction for all paragraphs
# ---------------------------
async def correct_paragraphs_async(
    docx_path, updated_docx_path, json_output_path, pdf_id="example_pdf_001"
):
    """
    Proofreads all paragraphs in a Word document asynchronously using GPT,
    updates the document, and writes a JSON report.
    Returns the number of improved paragraphs.
    """
    doc = Document(docx_path)
    total_word_changes = 0

    paragraphs_meta = []
    tasks = []

    para_id_counter = 1
    for paragraph in doc.paragraphs:
        original_text = paragraph.text.strip()

        if not original_text:
            # ✅ skip whitespace-only paragraphs entirely
            continue

        paragraphs_meta.append({
            "paragraph_id": para_id_counter,   # ✅ only count real paragraphs
            "original": paragraph.text,
            "proofread": paragraph.text,
            "original_token": [],
            "proofread_token": [],
            "original_text": [],
            "revised_text": [],
            "is_blank": False,
            "doc_paragraph": paragraph
        })

        print(f"Proofreading Paragraph {para_id_counter}: {paragraph.text}")
        tasks.append(async_gpt_proofread(para_id_counter, paragraph.text))
        para_id_counter += 1


    # Run GPT
    results = await asyncio.gather(*tasks)
    result_map = {pid: gpt_response for pid, _, gpt_response in results}

    for para in paragraphs_meta:
        pid = para["paragraph_id"]
        gpt_response = result_map.get(pid, {})

        corrected_text = gpt_response.get("corrected", para["original"])
        para["proofread"] = corrected_text
        para["original_token"] = gpt_response.get("original_token", [])
        para["proofread_token"] = gpt_response.get("proofread_token", [])
        para["original_text"] = [
            {
                "index": ch.get("original_idx"),
                "word": ch.get("original_word"),
                "type": "error",
            }
            for ch in gpt_response.get("changes", [])
            if ch.get("type") != "inserted" and ch.get("original_idx") is not None
        ]
        para["revised_text"] = [
            {
                "index": ch.get("proofread_idx"),
                "word": ch.get("proofread_word"),
                "type": ch.get("type"),
                # ✅ only keep real suggestions, avoid duplicates
                "suggestions": list(dict.fromkeys(
                    ch.get("suggestion", []) + [ch.get("proofread_word")]
                ))

            }
            for ch in gpt_response.get("changes", [])
            if ch.get("proofread_idx") is not None
        ]

        # Update Word document
        doc_paragraph = para["doc_paragraph"]
        if doc_paragraph.runs:
            ref_run = doc_paragraph.runs[0]
            doc_paragraph.clear()
            new_run = doc_paragraph.add_run(corrected_text)
            new_run.font.name = ref_run.font.name
            new_run.bold = ref_run.bold
            new_run.italic = ref_run.italic
            new_run.underline = ref_run.underline
            new_run.font.size = ref_run.font.size
            if ref_run.font.color and ref_run.font.color.rgb:
                new_run.font.color.rgb = ref_run.font.color.rgb
        else:
            doc_paragraph.text = corrected_text

        word_changes_count = sum(
            1 for ch in gpt_response.get("changes", [])
            if ch.get("type") in ["inserted", "changed", "replaced", "corrected"]
        )
        total_word_changes += word_changes_count

    # Save DOCX
    doc.save(updated_docx_path)
    print("Document updated with grammar corrections.")

    # ✅ Only keep non-blank paragraphs, and drop docx references
    clean_paragraphs = [
        {k: v for k, v in para.items() if k != "doc_paragraph"}
        for para in paragraphs_meta
        if not para["is_blank"]
    ]
    data = {"pdf_id": pdf_id, "paragraphs": clean_paragraphs}

    with open(json_output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print(f"Proofread JSON saved to {json_output_path}")

    return total_word_changes





# ---------------------------
# Update selected paragraphs in DOCX from JSON
# ---------------------------
def update_changes_on_pdf(final_pdf_path, updated_docx_path, json_output_path, paragraph_ids):
    with open(json_output_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    paragraphs_data = {str(p["paragraph_id"]): p for p in data.get("paragraphs", [])}
    paragraph_id_set = set(str(pid) for pid in paragraph_ids)

    doc = Document(updated_docx_path)
    updated_count = 0
    para_id_counter = 1  # ✅ real paragraph counter

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue  # ✅ skip blanks (keep ID sync with JSON)

        pid = str(para_id_counter)
        if pid in paragraph_id_set and pid in paragraphs_data:
            print(f"[UPDATE] Processing paragraph_id={pid}")
            proofread_text = paragraphs_data[pid].get("proofread", "")

            for change in paragraphs_data[pid].get("revised_text", []):
                if change["word"] == proofread_text:
                    # Mark that this change was applied but keep it modifiable
                    change["applied"] = True

            if paragraph.runs:
                ref_run = paragraph.runs[0]
                paragraph.clear()
                new_run = paragraph.add_run(proofread_text)
                new_run.font.name = ref_run.font.name
                new_run.bold = ref_run.bold
                new_run.italic = ref_run.italic
                new_run.underline = ref_run.underline
                new_run.font.size = ref_run.font.size
                if ref_run.font.color and ref_run.font.color.rgb:
                    new_run.font.color.rgb = ref_run.font.color.rgb
            else:
                paragraph.text = proofread_text

            updated_count += 1

        para_id_counter += 1  # ✅ increment only for non-blank

    doc.save(updated_docx_path)
    return updated_count



# ---------------------------
# FastAPI app and endpoint
# ---------------------------
app = FastAPI()

@app.get("/api/grammar-check")
async def grammar_check(
    mode: int = Query(...),
    file_code: str = Query(...),
    paragraph_id: str = Query("[]"),
):
    """
    mode=0: full run
    mode=1: apply changes only to selected paragraphs
    """

    # 1) Parse paragraph_id into a list[int]
    ids = [int(x) for x in re.findall(r"\d+", paragraph_id)]
    if mode == 1 and not ids:
        raise HTTPException(status_code=400, detail="No paragraph_id values provided.")

    start_time = time.time()

    # Paths
    pdf_path_original = os.path.abspath(f"original_pdfs/{file_code}.pdf")
    final_pdf_path    = os.path.abspath(f"processed_pdfs/{file_code}.pdf")
    docx_path         = os.path.abspath(f"parsing_words/{file_code}_temp.docx")
    updated_docx_path = os.path.abspath(f"parsing_words/{file_code}_updated.docx")
    json_output_path  = os.path.abspath(f"jsons/{file_code}.json")

    if mode == 0:
        # Full pipeline
        download_from_hostinger(file_code, pdf_path_original)        # <- make sure this pulls the ORIGINAL
        convert_pdf_to_word(pdf_path_original, docx_path)
        total_improvements = await correct_paragraphs_async(
            docx_path, updated_docx_path, json_output_path
        )
        convert_word_to_pdf(updated_docx_path, final_pdf_path)

    else:
        # 2) For partial updates, make sure we edit the LATEST processed assets
        # Add these helpers if you don’t have them yet:
        #   - download_processed_from_hostinger(file_code, dst_pdf)
        #   - download_json_from_hostinger(file_code, dst_json)
        download_processed_from_hostinger(file_code, final_pdf_path)
        download_json_from_hostinger(file_code, json_output_path)

        convert_pdf_to_word(final_pdf_path, docx_path)

        # 3) Pass the parsed list of ints, not the raw string
        total_improvements = update_changes_on_pdf(
            docx_path, updated_docx_path, json_output_path, ids
        )

        convert_word_to_pdf(updated_docx_path, final_pdf_path)

    # Upload latest artifacts back to Hostinger
    upload_processed_file(final_pdf_path, file_code)
    upload_parsing_words(updated_docx_path, file_code)
    upload_jsons(json_output_path, file_code)

    elapsed = time.time() - start_time

    # Option A (simplest): keep same filename but add a cache-buster token
    cache_bust = int(time.time())

    json_path = f"./jsons/{file_code}.json"

    if not os.path.exists(json_path):
        return {"error": "JSON not found"}

    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Convert param string "[1,3]" -> [1,3]
    paragraph_ids = []
    if paragraph_id:
        try:
            paragraph_ids = json.loads(paragraph_id)
        except Exception as e:
            print("⚠️ Failed to parse paragraph_id:", paragraph_id, e)

    # 🔍 Debug log
    print(f"[API] Mode={mode}, File={file_code}, Paragraphs={paragraph_ids}")

    # Iterate through data and log changes
    for para in data.get("paragraphs", []):
        if para.get("paragraph_id") in paragraph_ids:
            for change in para.get("revised_text", []):
                idx = change.get("index")
                new_word = change.get("word")

                # Find old word
                old_word = None
                for tok in para.get("proofread_token", []):
                    if tok["idx"] == idx:
                        old_word = tok["word"]
                        break

                print(
                    f"[API APPLY] Paragraph {para['paragraph_id']}, "
                    f"idx={idx}, old='{old_word}', new='{new_word}'"
                )

    return {
        "json_filename": os.path.basename(json_output_path),
        "final_pdf_filename": f"{os.path.basename(final_pdf_path)}?v={cache_bust}",
        "total_improvements": total_improvements,
        "elapsed_time_seconds": round(elapsed, 2),
    }


if __name__ == "__main__":
    # Run the FastAPI app with Uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)